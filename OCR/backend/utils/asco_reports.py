import os

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

FONT_DIR = os.path.join(os.path.dirname(__file__), "..", "static", "fonts")
pdfmetrics.registerFont(TTFont("DejaVu", os.path.join(FONT_DIR, "DejaVuSans.ttf")))

def generate_docx_from_asco(asco_data: dict, report_path: str, report_url: str):
    doc = Document()

    # ---- Title ----
    title = doc.add_heading("INSPECTION CERTIFICATE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Fields table ----
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Key"
    hdr_cells[1].text = "Value"

    for item in asco_data["fields"]:
        row = table.add_row().cells
        row[0].text = item["key"]
        row[1].text = item["value"]

    doc.add_paragraph()

    texts = asco_data["texts"]

    # ---- Inspection note ----
    p = doc.add_paragraph(texts["inspection_note"])
    p.runs[0].italic = True

    # ---- Conclusion ----
    doc.add_paragraph("Conclusion:", style="List Bullet")
    doc.add_paragraph(texts["conclusion"])

    # ---- Footer / Remarks ----
    center = doc.add_paragraph("ARIA SINA CONTROL INTERNATIONAL INSPECTION CO. (ASCO)")
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center.runs[0].bold = True

    remarks = doc.add_paragraph(texts["remarks"])
    remarks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remarks.runs[0].font.size = Pt(9)

    footer = doc.add_paragraph(texts["footer_note"])
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].bold = True

    doc.save(report_path)

    return report_url

def generate_pdf_from_asco(asco_data: dict, report_path: str, report_url: str):

    doc = SimpleDocTemplate(
        report_path,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    styles = getSampleStyleSheet()
    cell_style = styles["BodyText"]
    cell_style.wordWrap = "CJK"
    cell_style.fontName = "DejaVu"
    elements = []

    # ---- Title ----
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        alignment=TA_CENTER
    )
    elements.append(Paragraph("INSPECTION CERTIFICATE", title_style))
    elements.append(Spacer(1, 16))

    # ---- Fields table ----
    table_data = [
        [
            Paragraph("<b>Key</b>", styles["BodyText"]),
            Paragraph("<b>Value</b>", styles["BodyText"]),
        ]
    ]

    for item in asco_data["fields"]:
        table_data.append([
            Paragraph(item["key"], cell_style),
            Paragraph(item["value"], cell_style),
        ])

    table = Table(table_data, colWidths=[180, 320])
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 20))

    texts = asco_data["texts"]

    # ---- Inspection note (italic) ----
    elements.append(
        Paragraph(
            texts["inspection_note"],
            ParagraphStyle(
                "InspectionNote",
                parent=styles["BodyText"],
                fontName="Helvetica-Oblique",
                textColor=colors.grey,
            ),
        )
    )
    elements.append(Spacer(1, 14))

    # ---- Conclusion ----
    elements.append(Paragraph("<b>Conclusion:</b>", styles["BodyText"]))
    elements.append(Paragraph(texts["conclusion"], styles["BodyText"]))
    elements.append(Spacer(1, 16))

    # ---- Footer / Remarks ----
    elements.append(Paragraph(
        "<b>ARIA SINA CONTROL INTERNATIONAL INSPECTION CO. (ASCO)</b>",
        ParagraphStyle("CenterBold", parent=styles["BodyText"], alignment=TA_CENTER),
    ))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph(
        texts["remarks"],
        ParagraphStyle(
            "Remarks",
            parent=styles["BodyText"],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER,
        ),
    ))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph(
        texts["footer_note"],
        ParagraphStyle(
            "Footer",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            alignment=TA_CENTER,
        ),
    ))

    doc.build(elements)

    return report_url
